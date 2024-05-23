#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed May 15 16:22:54 2024

@author: ashutoshgoenka
"""
import pandas as pd
import numpy as np
import streamlit as st
from PIL import Image
import PIL
import os
import zipfile
from zipfile import ZipFile, ZIP_DEFLATED
import pathlib
import shutil
import docx
import docxtpl
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Inches


st.set_page_config(layout="wide")

try:
    shutil.rmtree("images_comp")
except:
    pass

name_index_dict= {}
image_size_dict = {}
original_image_size= {}
new_width_dict = {}
new_height_dict = {}
section_selected = ""

def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width


#Create a document with Landscape and saved
document = Document()
section = document.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height
document.save("Table_Word.docx")
# Document Created


# Upload Observation File
st.title("Upload Observation")
obs_file = st.file_uploader("Upload Observation Excel File", type=['csv','xlsx'],accept_multiple_files=False,key="fileUploader")
if obs_file is not None:
    df = pd.read_excel(obs_file)
    st.write(df)

    section_selected = df["Section"][0]
    
    df = df[df["Section"] == section_selected]
    
    df_rem = pd.read_excel("remedy_excel.xlsx")
    
    remedy_dict = {}
    for idx, val in df_rem.iterrows():
    #     print(val)
        remedy_dict[val["Observations"]] = val["Remedy"]
        
    df = df.dropna(thresh=5)
    
    df["Photo Start"] = pd.to_numeric(df['Photo Start']).astype('Int64')
    df["Photo End"] = df["Photo End"].fillna(0)
    df["Photo End"] = pd.to_numeric(df['Photo End']).astype('Int64')
    
    
    # Adding Photo column which specifies photo number
    order_of_section = []
    df["Photos"] = ""
    df["Action Needed"] = ""
    df["Observations + Location"] = ""
    
    for idx, val in df.iterrows():
    #     print(idx, val)
        if val["Section"] not in order_of_section:
            order_of_section.append(val["Section"])
        
        if np.isnan(val["Photo Start"]):
            pass
        elif val["Photo End"] == 0:
            df.loc[idx,"Photos"] = "Image " + str(val["Photo Start"])
        else:
            df.loc[idx,"Photos"] = "Image " + str(val["Photo Start"]) + " - " + str(val["Photo End"])
        
        df.loc[idx,"Action Needed"] = remedy_dict[val["Observations"]]
        
        val_temp = val["Observations"] +"\n"+"\n"+  val["Location"] + "\n"
        df.loc[idx,"Observations + Location"] = val_temp
    # Photo Column added
    
    col = ["Item", "Observations + Location", "Action Needed", "Category", "Photos", "Remarks/Action By"]
    df_final = df.copy(deep = True)
    df_final =  df_final.reindex(columns=col)
    
    st.write("Table to be added")
    st.write(df_final)
    #Add text table to word
    doc = docx.Document('Table_Word.docx')
    doc.add_heading(section_selected, 1)

    document.add_paragraph(section_selected)
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    t = doc.add_table(df_final.shape[0]+1, df_final.shape[1])
    t.style = 'Table Grid'
    t.allow_autofit = False
    t.columns[1].width = Cm(7.5)
    st.write(df_final.shape)
    # add the header rows.
    for j in range(df_final.shape[-1]):
        t.cell(0,j).text = df_final.columns[j]
    
    # add the rest of the data frame
    for i in range(df_final.shape[0]):
        for j in range(df_final.shape[-1]):
            t.cell(i+1,j).text = str(df_final.values[i,j])
    
    set_column_width(t.columns[1], docx.shared.Cm(7.5))
    set_column_width(t.columns[2], docx.shared.Cm(5.5))
    set_column_width(t.columns[3], docx.shared.Cm(2))
    # save the doc
    document.add_paragraph('')
    doc.save('Table_Word.docx')
    
    
    
    
    #Text Table added




#upload Images
st.title("Resize Images")
# st.write('My first app Hello *world!*')
up_files = st.file_uploader("Upload Image Files", type = ["png", "jpeg", "jpg"] ,accept_multiple_files=True)
# st.write(up_files)

def resize(img, new_width):
    width, height  = img.size
    ratio = height/width
    new_height = int(ratio*new_width)
    resized_image = img.resize((new_width, new_height), resample=PIL.Image.LANCZOS)
    return resized_image


def updateTable():
    # global up_files
    global folder
    document = Document("Table_Word.docx")
    document.add_heading(section_selected + " - Images", 2)
    
    table = document.add_table(rows = 7 , cols = 3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Item'     
    # hdr_cells[1].text = 'quantity'
    document.save("Table_Word.docx")
    counter = 0
    counter_cols = 0
    for file in folder.iterdir():
        cell = table.rows[counter].cells[counter_cols]
        cell._element.clear_content()
        picture = cell.add_paragraph().add_run().add_picture('images_comp/'+file.name, width=Inches(2.6))
        cell = table.rows[counter+1].cells[counter_cols]
        name = os.path.splitext(file.name)[0]
        cell.text = name
        if counter_cols<2:
            counter_cols = counter_cols + 1
        else:
            counter_cols = 0
            counter = counter+2
    document.save("Table_Word.docx")
    
    
def update_col():
    pass
    
                
def resize_image(img, width, height):

    # box = (size_to_scale, size_to_scale, size_to_scale, size_to_scale)
    box=((img.width-width)/2,((img.height-height)/2),(img.width+width)/2,((img.height+height)/2))
    im_resized = im.crop(box)
    return im_resized

title = st.text_input("Image Number to Start with", 1)
st.write("The image numbering will start from: ", title)

try:
    os.mkdir("images_comp")
except:
    pass

name_list = []

for i in range(len(up_files)):
    name_list.append("Image "+str(int(title)+i))



st.write(len(up_files))


for file in up_files:
    try:
        a = name_index_dict[file.name]
    except:
        name_index_dict[file.name] = 0

    
    # files = os.listdir("images")
    extensions = ["jpg", "jpeg", "png", "gif", "webp"]
    im = Image.open(file)
    ext = file.name.split(".")[-1]
    
    
    
    # Displaying Image
    im_width, im_height = im.size 
    original_image_size[file.name] = [im_width, im_height]
    try:
        b = image_size_dict[file.name]
    except:
        image_size_dict[file.name] = [im_width, im_height]
        
    
        
    
    st.write(im_width, im_height)
    size_to_scale = min(im_width,im_height)
    st.write(size_to_scale)
    # box = (size_to_scale, size_to_scale, size_to_scale, size_to_scale)
    box=((im_width-size_to_scale)/2,((im_height-size_to_scale)/2),(im_width+size_to_scale)/2,((im_height+size_to_scale)/2))
    im_resized = im.crop(box)
    
    try:
        a= next_width_dict[file.name]
        
    except:
        new_width_dict[file.name] = size_to_scale
        new_height_dict[file.name] = size_to_scale
    
    
    col1, col2, col3  = st.columns(3)
    with col1:
        st.image(file, width=350)
        # oi_width  = st.number_input("width", value = im_width)
        # oi_height = st.number_input("height", value = im_height)
        st.write(im_width, im_height)
        
    
    
    with col2:
        try:
            im_resized = resize_image(im, new_width_dict[file.name], new_height_dict[file.name])
        except:
            im_resized = im_resized
            
        
        col2_img = st.image(im_resized, width=350)
        st.write(im_resized.size)
        new_width_dict[file.name]  = st.number_input("new width", value = im_resized.size[0])
        new_height_dict[file.name] = st.number_input("new height", value = im_resized.size[1])
        # st.write(im_width, im_height)
        st.write(im_resized.size)
    
    
    with col3:
        try:
            im_resized_final = resize_image(im, new_width_dict[file.name], new_height_dict[file.name])
        except:
            im_resized_final = im_resized
        
        
        col3_img = st.image(im_resized_final, width=350)
        st.write(im_resized_final.size)
    
    # if im_resized.size[0] != new_width_dict[file.name] | im_resized.size[1] !=new_height_dict[file.name]:
    #     try:
    #         im_resized_final = resize_image(im, new_width_dict[file.name], new_height_dict[file.name])
    #     except:
    #         im_resized_final = im_resized
        
        
    #         col3_img = st.image(im_resized_final, width=350)
            
    
    
    
        
    # if im_resized.size[0] != new_width_dict[file.name] | im_resized.size[1] !=new_height_dict[file.name]:
    #     try:
    #         im_resized_final = resize_image(im, new_width_dict[file.name], new_height_dict[file.name])
    #     except:
    #         im_resized_final = im_resized
        
    #     with col3:
    #         col3_img = st.image(im_resized_final, width=350)
            
        
    
    
    
    
    # st.write(file)
    # st.image([file, im_resized], width=350)
    # st.image(im_resized, width = 250)
    name = os.path.splitext(file.name)[0]
    option =   st.selectbox(
            "File Name",
            tuple([name] + name_list),
            index=name_index_dict[file.name],
            )

    st.write("You selected:", option)
    list_temp = [name] + name_list
    position = list_temp.index(option)

    name_index_dict[name] = position
    # st.write(im.size)
    # im_resized = resize(im, 1000)
    # st.write(im_resized.size)
    im_width, im_height = im.size 
    st.write(im_width, im_height)
    size_to_scale = min(im_width,im_height)
    st.write(size_to_scale)
    # box = (size_to_scale, size_to_scale, size_to_scale, size_to_scale)
    box=((im_width-size_to_scale)/2,((im_height-size_to_scale)/2),(im_width+size_to_scale)/2,((im_height+size_to_scale)/2))
    im_resized = im.crop(box)
    # st.image(im_resized)    
    im_resized_final.save("images_comp/"+option+"."+ext)
    
    

zip_path = "images_compressed.zip"
directory_to_zip = "images_comp"
folder = pathlib.Path(directory_to_zip)
# st.write(folder)

# #Create a document with Landscape and saved
# document = Document()
# section = document.sections[0]
# section.orientation = WD_ORIENT.LANDSCAPE
# new_width, new_height = section.page_height, section.page_width
# section.page_width = new_width
# section.page_height = new_height
# document.save("Table_Word.docx")
# # Document Created


with ZipFile(zip_path, 'w', ZIP_DEFLATED) as zip:
    for file in folder.iterdir():
        zip.write(file, arcname=file.name)
        
with open("images_compressed.zip", "rb") as fp:
    btn = st.download_button(
        label="Download ZIP",
        data=fp,
        file_name="images_compressed.zip",
        mime="application/zip"
    )

updateTable()

with open("Table_Word.docx", "rb") as fp:

    btn_1 = st.download_button(
            label="Word File with Table",
            data=fp,
            file_name="Table_Word_docx",
            mime="docx"
            )
    
    # st.write(btn_1)
    
    # if btn_1:
    #     st.write("Running Update Function")
    #     updateTable(up_files)

os.remove(zip_path)
shutil.rmtree("images_comp")

# st.download_button("Download Images", file_name="bali.jpeg")
    
    # for file in files:
    # ext = im.name.split(".")[-1]
    # if ext in extensions:
    #     # im = Image.open("images/"+file)
        
        
    #     im_resized = resize(im, 400)
    #     filepath = "images/"+file+".jpg"
    #     im_resized.save(filepath)
        
